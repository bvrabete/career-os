import unittest
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import RGBColor

from docx_generator import (
    generate_docx,
    add_hyperlink,
    add_formatted_runs,
    _process_part,
    _add_title_header,
    _add_section_header,
    _add_subsection_header,
    _add_role_header,
    _add_horizontal_rule,
    _add_bullet_item,
    _add_standard_paragraph
)


class TestDocxGenerator(unittest.TestCase):

    def setUp(self) -> None:
        self.doc = Document()
        self.color_primary = RGBColor(15, 34, 64)
        self.color_secondary = RGBColor(80, 80, 80)
        self.color_text = RGBColor(30, 30, 30)

    def test_add_hyperlink(self) -> None:
        p = self.doc.add_paragraph()
        add_hyperlink(p, "https://example.com", "Example Link")
        self.assertTrue(len(p.runs) == 0) # Hyperlink run is appended under OxmlElement child nodes
        self.assertIn("Example Link", p._p.xml)

    def test_process_part_styles(self) -> None:
        p = self.doc.add_paragraph()
        _process_part(p, "***bold italic***")
        self.assertEqual(p.runs[0].text, "bold italic")
        self.assertTrue(p.runs[0].bold)
        self.assertTrue(p.runs[0].italic)

        p2 = self.doc.add_paragraph()
        _process_part(p2, "**bold**")
        self.assertEqual(p2.runs[0].text, "bold")
        self.assertTrue(p2.runs[0].bold)

        p3 = self.doc.add_paragraph()
        _process_part(p3, "*italic*")
        self.assertEqual(p3.runs[0].text, "italic")
        self.assertTrue(p3.runs[0].italic)

        p4 = self.doc.add_paragraph()
        _process_part(p4, "[link](https://test.org)")
        self.assertIn("link", p4._p.xml)

        p5 = self.doc.add_paragraph()
        _process_part(p5, "[broken_link")
        self.assertEqual(p5.runs[0].text, "[broken_link")

    def test_add_formatted_runs(self) -> None:
        p = self.doc.add_paragraph()
        add_formatted_runs(p, "Standard text **bold** and *italic* and [link](https://test.org).")
        self.assertTrue(len(p.runs) >= 4)

    def test_add_headers(self) -> None:
        _add_title_header(self.doc, "Main Title", self.color_primary)
        _add_section_header(self.doc, "Section Header", self.color_primary)
        _add_subsection_header(self.doc, "Subsection Header", self.color_primary)
        _add_role_header(self.doc, "Role Header", self.color_primary)
        
        self.assertEqual(self.doc.paragraphs[0].text, "Main Title")
        self.assertEqual(self.doc.paragraphs[1].text, "Section Header")
        self.assertEqual(self.doc.paragraphs[2].text, "Subsection Header")
        self.assertEqual(self.doc.paragraphs[3].text, "Role Header")

    def test_add_horizontal_rule(self) -> None:
        _add_horizontal_rule(self.doc)
        self.assertIn("_____", self.doc.paragraphs[0].text)

    def test_add_bullet_item(self) -> None:
        _add_bullet_item(self.doc, "Bullet text **bold**", self.color_text)
        p = self.doc.paragraphs[0]
        style_name = getattr(p.style, "name", None)
        self.assertEqual(style_name, "List Bullet")

    def test_add_standard_paragraph(self) -> None:
        _add_standard_paragraph(self.doc, "Standard text line", 1, self.color_secondary, self.color_text)
        self.assertEqual(self.doc.paragraphs[0].text, "Standard text line")

    def test_add_standard_paragraph_contact(self) -> None:
        _add_standard_paragraph(self.doc, "London, UK | +44 | email@address.com", 2, self.color_secondary, self.color_text)
        p = self.doc.paragraphs[0]
        # Should center align contact-like lines at indices under 8
        self.assertEqual(p.alignment, 1) # WD_ALIGN_PARAGRAPH.CENTER is 1

    def test_generate_docx_success(self) -> None:
        test_md = """# Title
London, UK | +44 | email@address.com

---

## Experience
### Job Title - Company (Present)
- Bullets go here **bold**
- Bullet 2 *italic* and [link](https://google.com)

#### Subrole Header
Paragraph here."""

        with tempfile.TemporaryDirectory() as tmp_dir:
            out_path = Path(tmp_dir) / "output.docx"
            res = generate_docx(test_md, str(out_path))
            self.assertTrue(res)
            self.assertTrue(out_path.exists())

    def test_generate_docx_failure(self) -> None:
        # Invalid output location (e.g. writing to empty string filename or locked path)
        res = generate_docx("# Simple MD", "")
        self.assertFalse(res)


if __name__ == "__main__":
    unittest.main()
