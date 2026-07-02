"""
Module for generating Word Documents (.docx) from Markdown content.
Uses pure-Python python-docx library to ensure 100% OS-independence.
"""

import logging
import re
from pathlib import Path
from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)


def add_hyperlink(paragraph: Any, url: str, text: str) -> Any:
    """
    Helper function for adding XML hyperlinks (clickable links) to a paragraph
    using standard RELATIONSHIP_TYPE in python-docx.
    """
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')

    # Premium styling: Medium blue color + single underline
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0056B3')
    r_pr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    r_pr.append(u)

    new_run.append(r_pr)

    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _process_part(paragraph: Any, part: str) -> None:
    """Applies inline formatting styles to a specific token part."""
    if part.startswith('***') and part.endswith('***'):
        run = paragraph.add_run(part[3:-3])
        run.bold = True
        run.italic = True
    elif part.startswith('**') and part.endswith('**'):
        run = paragraph.add_run(part[2:-2])
        run.bold = True
    elif part.startswith('*') and part.endswith('*'):
        run = paragraph.add_run(part[1:-1])
        run.italic = True
    elif part.startswith('[') and ']' in part and part.endswith(')'):
        match = re.match(r'\[([^\]]*)\]\(([^\)]*)\)', part)
        if match:
            link_text, url = match.groups()
            add_hyperlink(paragraph, url, link_text)
        else:
            paragraph.add_run(part)
    else:
        paragraph.add_run(part)


def add_formatted_runs(paragraph: Any, text: str) -> None:
    """
    Parses and appends inline formatting (bold, italic, links) to a paragraph.
    """
    # Tokenize by bold, italic, and links using non-backtracking negated character classes
    pattern = re.compile(
        r'(\*\*\*[^\*]*\*\*\*|\*\*[^\*]*\*\*|\*[^\*]*\*|\[[^\]]*\]\([^\)]*\))')
    parts = pattern.split(text)

    for part in parts:
        if part:
            _process_part(paragraph, part)


def _add_title_header(doc: Any, text: str, color_primary: RGBColor) -> None:
    """Formats and adds the main title to the document."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)

    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(22)
    run.font.color.rgb = color_primary


def _add_section_header(doc: Any, text: str, color_primary: RGBColor) -> None:
    """Formats and adds a section header (Heading 2) to the document."""
    p = doc.add_heading(level=2)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0)

    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.color.rgb = color_primary


def _add_role_header(doc: Any, text: str, color_primary: RGBColor) -> None:
    """Formats and adds a sub-header or role header to the document."""
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)

    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = color_primary


def _add_horizontal_rule(doc: Any) -> None:
    """Adds an elegant, centered horizontal rule separator."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(1)

    run = p.add_run("______________________________________________________________________")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(210, 210, 210)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_bullet_item(doc: Any, text: str, color_text: RGBColor) -> None:
    """Formats and adds a single list bullet item to the document."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15

    add_formatted_runs(p, text)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        if not run.font.color.rgb:
            run.font.color.rgb = color_text


def _add_standard_paragraph(
    doc: Any, line: str, idx: int, color_secondary: RGBColor, color_text: RGBColor
) -> None:
    """Formats and adds a standard body paragraph or contact info line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15

    is_contact_line = any(
        kw in line.lower() for kw in ["@ ", "phone", "linkedin", "address", "email", "github"]
    ) or "|" in line or "  " in line

    if is_contact_line and idx < 8:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)

    add_formatted_runs(p, line)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(9.5) if is_contact_line else Pt(10.5)
        if is_contact_line:
            run.font.color.rgb = color_secondary
        elif not run.font.color.rgb:
            run.font.color.rgb = color_text


from utils import clean_markdown_wrapper as _clean_markdown_wrapper


def generate_docx(md_content: str, output_path: str) -> bool:
    """
    Converts Markdown content to a Word Document (.docx) using pure-Python python-docx library,
    ensuring 100% OS-independence (no Pandoc system binary required).
    """
    logger.info(f"Generating pure-Python DOCX for {output_path}...")
    try:
        from docx import Document
        doc = Document()

        # Set standard elegant margins (0.75 inches all around)
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Theme color tokens (Modern Slate Blue theme)
        color_primary = RGBColor(15, 34, 64)       # Deep slate blue
        color_secondary = RGBColor(80, 80, 80)     # Muted grey
        color_text = RGBColor(30, 30, 30)          # Charcoal dark body

        # Clean leading/trailing markdown code blocks if the entire content is wrapped
        cleaned_md = _clean_markdown_wrapper(md_content)
        lines = cleaned_md.splitlines()

        idx = 0
        while idx < len(lines):
            line = lines[idx].strip()

            if not line:
                idx += 1
                continue

            if line.startswith("# "):
                _add_title_header(doc, line[2:].strip(), color_primary)
            elif line.startswith("## "):
                _add_section_header(doc, line[3:].strip(), color_primary)
            elif line.startswith("### "):
                _add_role_header(doc, line[4:].strip(), color_primary)
            elif line.startswith("---") or (line.startswith("***") and len(line) >= 3 and not line.strip("* -")):
                _add_horizontal_rule(doc)
            elif line.startswith(("- ", "* ", "+ ")):
                _add_bullet_item(doc, line[2:].strip(), color_text)
            else:
                _add_standard_paragraph(doc, line, idx, color_secondary, color_text)

            idx += 1

        # Save to target location
        out_path = Path(output_path)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))
        logger.info(f"Successfully saved pure-Python DOCX to {output_path}")
        return True

    except Exception as e:
        logger.exception(f"Failed to generate DOCX via python-docx: {str(e)}")
        return False


if __name__ == "__main__":
    # Self-test block
    logging.basicConfig(level=logging.INFO,
                        format='%(asctime)s - %(levelname)s - %(message)s')
    test_md = """# Joanna Doe, Great Comp Ltd
London, UK | +44 123 4567 | [LinkedIn](https://linkedin.com/in/ja-doe) | [Email](mailto:joana.doe@example.com)

---

## Work Experience

### Lead Delivery Coordinator - O My Ltd  (2017 - Present)
- Led the delivery orchestration and dispatch operations of agentic logistics.
- Optimized and saved **15% on logistics costs** in EMEA.
- Wrote **Python scripts** to parse routing databases and generate schedules.

### Senior Support Specialist - Xoft (2011 - 2016)
- Managed client bookkeeping and pilot platform presentation materials."""
    generate_docx(test_md, "test_pure_python_cv.docx")
